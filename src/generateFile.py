import random, string
import shutil
import datetime
import hashlib
from docx.oxml import OxmlElement
from docx import Document

def generateFile():
	source_file = open('Document.docx','rb')
	data_dump=random_string(random.randint(100,500))
	fileId=hashlib.sha256(str(datetime.datetime.now()).encode('utf-8')).hexdigest()
	filepath='output/Document_'+fileId+'.docx'
	destination_file=open(filepath,'wb')
	shutil.copyfileobj(source_file,destination_file)
	generateHiddenParagraph(filepath,data_dump)
	return filepath
	
	
def random_string(length):
	letters=string.ascii_lowercase
	return ''.join(random.choice(letters) for i in range(length))

def generateHiddenParagraph(filepath,data_dump):
	doc=Document(filepath)
	p=doc.add_paragraph()
	setHiddenProperty(p)
	r=p.add_run()
	r.text=data_dump
	r.font.hidden=True
	doc.save(filepath)
	
def setHiddenProperty(p):
	pPr=OxmlElement('w:pPr')
	rPr=OxmlElement('w:rPr')
	v=OxmlElement('w:vanish')
	rPr.append(v)
	pPr.append(rPr)
	p._p.append(pPr)
	
	
